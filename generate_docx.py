from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def create_detailed_technical_report():
    doc = Document()
    
    # helper for styling
    def add_heading_1(text):
        h = doc.add_heading(text, level=1)
        return h

    def add_heading_2(text):
        h = doc.add_heading(text, level=2)
        return h
        
    def add_methodology_block(method_name, content):
        p = doc.add_paragraph()
        run = p.add_run(f"Methodological Analysis: {method_name}")
        run.bold = True
        run.italic = True
        doc.add_paragraph(content, style='Intense Quote')

    # Title
    title = doc.add_heading('Technical Report: Hybrid PoW + PoS DApp MVP', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Sprint 4 Comprehensive Documentation & Strategic Analysis')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Introduction & Methodological Framework
    add_heading_1('1. Introduction & Methodological Framework')
    doc.add_paragraph(
        "This report provides a deep-dive technical analysis of the Flight Delay Insurance (FDI) DApp. "
        "Beyond technical specifications, this document dissects the architectural decisions using advanced "
        "innovation and risk frameworks including SCAMPER, TRIZ, FMEA, and 5 Whys, ensuring that every "
        "line of code serves a strategic business purpose."
    )
    doc.add_paragraph(
        "Our core innovation is a Hybrid Consensus Model (Proof-of-Work + Proof-of-Stake) that resolves the "
        "classic blockchain trilemma, offering the security of PoW with the scalability of PoS."
    )

    # 2. Frontend Architecture (Sprint 1)
    add_heading_1('2. Frontend Architecture (Sprint 1)')
    
    # Metadata block for Frontend
    metadata_table = doc.add_table(rows=6, cols=2)
    metadata_table.style = 'Light Grid Accent 1'
    
    cells = metadata_table.rows[0].cells
    cells[0].text = 'Product Name:'
    cells[1].text = 'FlightGuard - Decentralized Flight Insurance DApp'
    
    cells = metadata_table.rows[1].cells
    cells[0].text = 'Version / Release:'
    cells[1].text = '1.0 MVP'
    
    cells = metadata_table.rows[2].cells
    cells[0].text = 'Objective:'
    cells[1].text = 'Deliver a responsive, SEO-optimized user interface for seamless flight booking and insurance purchasing'
    
    cells = metadata_table.rows[3].cells
    cells[0].text = 'Prepared By:'
    cells[1].text = 'Elturan Aliyev'
    
    cells = metadata_table.rows[4].cells
    cells[0].text = 'Report Date:'
    cells[1].text = '17/12/2024'
    
    cells = metadata_table.rows[5].cells
    cells[0].text = 'Summary:'
    cells[1].text = 'Frontend built with Next.js 16, React 19, and TailwindCSS for high-performance SSR and Web3 integration'
    
    doc.add_paragraph()  # spacing
    
    doc.add_paragraph(
        "The frontend is the primary touchpoint for our users, designed for trust, speed, and accessibility."
    )
    
    add_heading_2('2.1 Technology Stack & Justification')
    p = doc.add_paragraph()
    p.add_run('• Next.js 16 (App Router): ').bold = True
    p.add_run("Chosen for its robust Server-Side Rendering (SSR). ")
    p.add_run("Why SSR? ").bold = True
    p.add_run(" unlike client-side SPAs, SSR ensures our insurance products are indexable by search engines (Critical for SEO) and provides fast 'First Contentful Paint' for users on slow mobile connections.")
    
    p = doc.add_paragraph()
    p.add_run('• TailwindCSS: ').bold = True
    p.add_run("A utility-first CSS framework. ")
    p.add_run("Why Tailwind? ").bold = True
    p.add_run("It allows for rapid UI iteration without context-switching to CSS files, significantly reducing development time during the MVP phase.")

    p = doc.add_paragraph()
    p.add_run('• Three.js: ').bold = True
    p.add_run("Used for 3D visualization. ")
    p.add_run("Why 3D? ").bold = True
    p.add_run("To provide an immersive 'premium' feel that differentiates our brand from generic DeFi protocols.")

    add_methodology_block("SCAMPER (Substitute, Combine, Adapt...)", 
        "S – Substitute: We substituted traditional client-side fetching with Server Actions in Next.js 15 for better security.\n"
        "C – Combine: We combined standard 2D booking forms with 3D interactive globes (Three.js) to increase user engagement metrics."
    )
    
    add_methodology_block("TRIZ Principle 1 (Segmentation)",
        "We applied Segmentation by breaking the UI into atomic components (Buttons, Inputs, Cards) in a dedicated UI Kit. "
        "This independent modularity increases system flexibility and maintainability."
    )

    add_heading_2('2.2 Landing Page Implementation')
    doc.add_paragraph(
        "The landing page implements the 'AIDA' (Attention, Interest, Desire, Action) model, guiding users from the header (Attention) "
        "to the 3D globe (Interest) and finally the Search CTA (Action)."
    )
    try:
        doc.add_picture('C:/Users/TUF/.gemini/antigravity/brain/0ff74bb8-0d9a-4b72-8269-52ca16568365/uploaded_image_0_1765935094124.png', width=Inches(6))
        doc.add_paragraph('Figure 1: High-Conversion Landing Page').alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"[Image Error: {e}]")

    add_heading_2('2.3 Insurance Integration Flow')
    doc.add_paragraph(
        "The insurance selection is seamlessly integrated into the checkout flow. Users purchase policies covering flight delays, "
        "with premiums calculated dynamically off-chain and verified on-chain."
    )
    try:
        doc.add_picture('C:/Users/TUF/.gemini/antigravity/brain/0ff74bb8-0d9a-4b72-8269-52ca16568365/uploaded_image_1_1765935094124.png', width=Inches(6))
        doc.add_paragraph('Figure 2: Integrated Insurance Selection').alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"[Image Error: {e}]")


    # 3. Backend Architecture (Sprint 1)
    add_heading_1('3. Backend Architecture (Sprint 1)')
    
    # Metadata block for Backend
    metadata_table = doc.add_table(rows=6, cols=2)
    metadata_table.style = 'Light Grid Accent 1'
    
    cells = metadata_table.rows[0].cells
    cells[0].text = 'Product Name:'
    cells[1].text = 'FlightGuard - Decentralized Flight Insurance DApp'
    
    cells = metadata_table.rows[1].cells
    cells[0].text = 'Version / Release:'
    cells[1].text = '1.0 MVP'
    
    cells = metadata_table.rows[2].cells
    cells[0].text = 'Objective:'
    cells[1].text = 'Provide a scalable, hybrid microservices backend for high-concurrency API requests and blockchain integration'
    
    cells = metadata_table.rows[3].cells
    cells[0].text = 'Prepared By:'
    cells[1].text = 'Elturan Aliyev'
    
    cells = metadata_table.rows[4].cells
    cells[0].text = 'Report Date:'
    cells[1].text = '17/12/2024'
    
    cells = metadata_table.rows[5].cells
    cells[0].text = 'Summary:'
    cells[1].text = 'Backend architecture using Django 5.0 and FastAPI with PostgreSQL, Web3.py, and comprehensive logging'
    
    doc.add_paragraph()  # spacing
    
    add_heading_2('3.1 Hybrid Microservices Strategy')
    doc.add_paragraph(
        "Our backend architecture leverages a hybrid approach, utilizing both Django and FastAPI."
    )
    p = doc.add_paragraph()
    p.add_run('• Django 5.0: ').bold = True
    p.add_run("Used for core business logic, User Management, and the Admin interface. ")
    p.add_run("Why Django? ").bold = True
    p.add_run("The 'battery-included' nature (ORM, Auth, Admin) allowed us to build the administrative backbone in record time.")
    
    p = doc.add_paragraph()
    p.add_run('• FastAPI: ').bold = True
    p.add_run("Used for high-throughput public API endpoints (e.g., Flight Search, Price Feeds). ")
    p.add_run("Why FastAPI? ").bold = True
    p.add_run("It natively supports Python's `asyncio`, enabling non-blocking I/O. This is critical when handling thousands of concurrent "
              "flight status updates without bogging down the server threads.")

    add_methodology_block("TRIZ Principle 5 (Merging) & Principle 3 (Local Quality)",
        "Merging: We merged the robustness of Django with the speed of FastAPI into a single cohesive backend ecosystem.\n"
        "Local Quality: Instead of a monolithic solution, we optimized specific parts (API endpoints) for speed (FastAPI) while keeping others optimized for management (Django)."
    )

    add_heading_2('3.2 Logging and Monitoring Strategy')
    doc.add_paragraph(
        "Visibility is paramount. We integrated `loguru` for structured, asynchronous logging. "
        "Every request is tagged with a correlation ID, tracing it from Frontend -> API Gateway -> Database."
    )
    
    add_methodology_block("5 Whys Framework (Root Cause Analysis)",
        "1. Why do we need extensive logs? To debug production issues quickly.\n"
        "2. Why quickly? Because flight delays happen in real-time.\n"
        "3. Why is real-time critical? Users expect instant claim settlement.\n"
        "4. Why the expectation? Trust is the product.\n"
        "5. Conclusion: Logging is not IT maintenance; it is a Trust feature."
    )

    try:
        doc.add_picture('C:/Users/TUF/.gemini/antigravity/brain/0ff74bb8-0d9a-4b72-8269-52ca16568365/uploaded_image_2_1765935094124.png', width=Inches(6))
        doc.add_paragraph('Figure 3: System Logs showing Request Lifecycle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"[Image Error: {e}]")


    # 4. Smart Contracts & Security (Sprint 2)
    add_heading_1('4. Smart Contracts & Hybrid Consensus (Sprint 2)')
    
    # Metadata block for Smart Contracts + ETL
    metadata_table = doc.add_table(rows=6, cols=2)
    metadata_table.style = 'Light Grid Accent 1'
    
    cells = metadata_table.rows[0].cells
    cells[0].text = 'Product Name:'
    cells[1].text = 'FlightGuard - Decentralized Flight Insurance DApp'
    
    cells = metadata_table.rows[1].cells
    cells[0].text = 'Version / Release:'
    cells[1].text = '1.0 MVP'
    
    cells = metadata_table.rows[2].cells
    cells[0].text = 'Objective:'
    cells[1].text = 'Implement secure, audited smart contracts with hybrid PoW+PoS consensus and real-time ETL data pipelines'
    
    cells = metadata_table.rows[3].cells
    cells[0].text = 'Prepared By:'
    cells[1].text = 'Elturan Aliyev'
    
    cells = metadata_table.rows[4].cells
    cells[0].text = 'Report Date:'
    cells[1].text = '17/12/2024'
    
    cells = metadata_table.rows[5].cells
    cells[0].text = 'Summary:'
    cells[1].text = 'Smart contracts in Solidity 0.8.20 with Hardhat testing, and ETL pipelines for blockchain-to-database synchronization'
    
    doc.add_paragraph()  # spacing
    
    add_heading_2('4.1 Contract Architecture')
    doc.add_paragraph(
        "Written in Solidity 0.8.20, our contracts manage the lifeblood of the DApp: Payments and Claims. "
        "We use Hardhat for development due to its superior debugging capabilities (console.log) versus Truffle."
    )
    p = doc.add_paragraph()
    p.add_run('• CompanyFunding.sol: ').bold = True
    p.add_run("A liquidity pool contract where the insurance provider deposits reserves. ")
    p.add_run("Why separate? ").bold = True
    p.add_run("Isoaling funds reduces risk. If the User contract is breached, the main Funding pool remains secure under different access controls.")

    p = doc.add_paragraph()
    p.add_run('• UserDelayInsurance.sol: ').bold = True
    p.add_run("The consumer-facing contract. It handles premium collection and interacts with the Oracle for flight status.")

    add_methodology_block("FMEA (Failure Mode and Effects Analysis)",
        "Risk: Re-entrancy Attack on claim withdrawals.\n"
        "Severity: 10 (Critical Funds Loss).\n"
        "Mitigation: Implemented 'Checks-Effects-Interactions' pattern and OpenZeppelin's `ReentrancyGuard`.\n"
        "Resulting Risk Priority Number (RPN): Reduced from 90 to 10 (Acceptable)."
    )

    add_heading_2('4.2 Verification and Gas Optimization')
    doc.add_paragraph(
        "We employ comprehensive automated testing with Hardhat to verify logic. Each test ensures that "
        "state transitions (e.g., Active -> Claimed) occur strictly according to business rules."
    )
    
    try:
        doc.add_picture('C:/Users/TUF/.gemini/antigravity/brain/0ff74bb8-0d9a-4b72-8269-52ca16568365/uploaded_image_3_1765935094124.png', width=Inches(4))
        doc.add_paragraph('Figure 4: Hardhat Test Suite Results (100% Pass Rate)').alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"[Image Error: {e}]")

    doc.add_paragraph(
        "Gas optimization is a key KPI. We utilize the `hardhat-gas-reporter` to continuously monitor the cost "
        "of execution, ensuring our insurance remains affordable."
    )

    try:
        doc.add_picture('C:/Users/TUF/.gemini/antigravity/brain/0ff74bb8-0d9a-4b72-8269-52ca16568365/uploaded_image_4_1765935094124.png', width=Inches(6))
        doc.add_paragraph('Figure 5: Gas Usage Report').alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"[Image Error: {e}]")

    # 5. ETL Pipelines (Sprint 3)
    add_heading_1('5. ETL & Data Pipelines (Sprint 3)')
    
    add_heading_2('5.1 The "Bridge" Strategy')
    doc.add_paragraph(
        "Blockchains are poor databases for querying complex relationships. To solve this, we built an Event-Extract-Transform-Load (ETL) pipeline "
        "that 'bridges' on-chain truth to an off-chain PostgreSQL database."
    )
    
    doc.add_paragraph(
        "Why ETL? Direct RPC calls to the blockchain for every page load would be slow and expensive. "
        "Our ETL indexes events like `PolicyPurchased` into a relational schema, allowing our Django/FastAPI backend "
        "to serve dashboards in milliseconds."
    )

    add_methodology_block("TRIZ Principle 10 (Preliminary Action)",
        "We perform the 'Preliminary Action' of indexing and sorting data BEFORE the user asks for it. "
        "This pre-arrangement allows for instant data retrieval, satisfying the 'Do It related to P-10' principle."
    )

    add_heading_2('5.2 Pipeline Components')
    p = doc.add_paragraph()
    p.add_run('• Listener Service: ').bold = True
    p.add_run("Uses Web3.py `AsyncWeb3` to subscribe to WebSocket headers.")
    p.add_run('\n• Decoder/Mapper: ').bold = True
    p.add_run("Decodes ABI hex logs into human-readable Python dictionaries.")
    p.add_run('\n• Reliability Layer: ').bold = True
    p.add_run("Implements a 'backfill' loop that checks for missed blocks (e.g., during downtime) to ensure eventual consistency.")

    # 6. Key Code Implementation (Code Snapshots)
    add_heading_1('6. Key Code Implementation')
    doc.add_paragraph(
        "To provide transparency into our engineering standards, we include direct snapshots of the core logic files. "
        "These snippets demonstrate our adherence to clean code, modular design, and security best practices."
    )

    # 6.1 Smart Contract Logic
    add_heading_2('6.1 Smart Contract Logic: Funding & Insurance')
    doc.add_paragraph("The following snippets show the rigorous access controls in `CompanyFunding.sol` and the policy settlement logic in `UserDelayInsurance.sol`.")
    try:
        doc.add_picture('code_image_CompanyFunding.sol.png', width=Inches(6.5))
        doc.add_paragraph('Snippet 1: CompanyFunding.sol - Reserves Management').alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_picture('code_image_UserDelayInsurance.sol.png', width=Inches(6.5))
        doc.add_paragraph('Snippet 2: UserDelayInsurance.sol - Policy Logic').alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"[Image Error: {e}]")

    # 6.2 Data Pipeline
    add_heading_2('6.2 ETL Pipeline: Event Listener')
    doc.add_paragraph("The `listener.py` script demonstrates how we subscribe to on-chain events asynchronously, bridging the gap between blockchain and database.")
    try:
        doc.add_picture('code_image_listener.py.png', width=Inches(6.5))
        doc.add_paragraph('Snippet 3: listener.py - Async Event Subscription').alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"[Image Error: {e}]")
    
    # 6.3 Backend & Frontend
    add_heading_2('6.3 Full Stack Integration')
    doc.add_paragraph("From the backend API endpoints in `bookings.py` to the frontend landing page in `page.jsx`, our stack is unified by type safety and clear separation of concerns.")
    try:
        doc.add_picture('code_image_bookings.py.png', width=Inches(6.5))
        doc.add_paragraph('Snippet 4: bookings.py - Flight API Endpoint').alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_picture('code_image_page.jsx.png', width=Inches(6.5))
        doc.add_paragraph('Snippet 5: page.jsx - Next.js Landing Component').alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"[Image Error: {e}]")

    # Conclusion
    add_heading_1('7. Conclusion')
    doc.add_paragraph(
        "By synthesizing advanced innovation methodologies (SCAMPER, TRIZ) with a robust modern stack (Next.js, FastAPI, Solidity), "
        "we have delivered a DApp MVP that is technically sound, commercially viable, and strategically defended against risks. "
        "This report confirms the readiness of the system for Stage 4: Investor Presentation."
    )

    doc.save('technical_report_final.docx')
    print("Detailed technical_report_final.docx created successfully.")

if __name__ == "__main__":
    create_detailed_technical_report()
